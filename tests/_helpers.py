"""测试 fixture 辅助：构造特殊 docx 用于边界场景测试。"""
import os
import zipfile
from pathlib import Path

from docx import Document


def make_garbage_docx(path):
    """写入垃圾内容到 .docx 文件（不是 zip）。"""
    with open(path, "wb") as f:
        f.write(b"this is not a real docx, just garbage bytes")


def make_corrupt_docx(path):
    """截断正常 docx 的后半部分（zip 中央目录损坏）。"""
    src = Path(__file__).resolve().parent / "_corrupt_source.docx"
    if not src.exists():
        doc = Document()
        doc.add_paragraph("test")
        doc.save(str(src))
    with open(src, "rb") as f:
        data = f.read()
    with open(path, "wb") as f:
        f.write(data[: max(100, len(data) // 3)])


def make_docx_with_nested_table(path):
    """构造含嵌套表的 docx。"""
    doc = Document()
    doc.add_paragraph("外层段落")
    tbl = doc.add_table(rows=2, cols=2)
    inner_cell = tbl.rows[0].cells[0]
    inner_cell.add_paragraph("嵌套前")
    inner_tbl = inner_cell.add_table(rows=1, cols=1)
    inner_tbl.rows[0].cells[0].text = "嵌套内容"
    inner_cell.add_paragraph("嵌套后")
    doc.save(path)


def make_docx_with_empty_table(path):
    """构造含空单元格的 docx（1 行 1 列但无文字）。"""
    doc = Document()
    doc.add_paragraph("前")
    tbl = doc.add_table(rows=1, cols=1)
    doc.add_paragraph("后")
    doc.save(path)


def make_docx_with_unrecognized_paragraph(path):
    """构造含无法识别段落的 docx。"""
    doc = Document()
    doc.add_paragraph("正常正文段落")
    doc.add_paragraph("12345")  # 纯数字，不匹配任何标题/正文模式
    doc.save(path)
