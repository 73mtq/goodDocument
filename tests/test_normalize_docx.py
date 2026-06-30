import json
import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]


def load_config():
    return json.loads((ROOT / "config.json").read_text(encoding="utf-8"))


class NormalizeDocxTests(unittest.TestCase):
    def test_inspect_docx_reports_structure(self):
        from normalizer import inspect_docx

        source = ROOT / "规范文档示例.docx"

        info = inspect_docx(source)

        self.assertGreater(info.paragraph_count, 0)
        self.assertEqual(info.table_count, 1)
        self.assertEqual(info.image_count, 1)
        self.assertIn(1, info.heading_levels)
        self.assertIn(2, info.heading_levels)
        self.assertTrue(info.has_header_footer)
        self.assertFalse(info.has_complex_objects)

    def test_normalize_docx_preserves_structure_and_applies_core_formatting(self):
        from normalizer import normalize_docx

        cfg = load_config()
        source = ROOT / "规范文档示例.docx"
        with tempfile.TemporaryDirectory() as tmp:
            output = Path(tmp) / "规范文档示例_规范化.docx"

            original = Document(source)
            normalize_docx(cfg, source, output)
            normalized = Document(output)

            self.assertTrue(output.exists())
            self.assertNotEqual(source.resolve(), output.resolve())
            self.assertEqual(len(original.paragraphs), len(normalized.paragraphs))
            self.assertEqual(len(original.tables), len(normalized.tables))
            self.assertEqual(original.paragraphs[0].text, normalized.paragraphs[0].text)

            margins = cfg["page"]["margins"]
            section = normalized.sections[0]
            self.assertAlmostEqual(section.top_margin.cm, margins["top"], places=1)
            self.assertAlmostEqual(section.left_margin.cm, margins["left"], places=1)
            self.assertIn("PAGE", section.footer._element.xml)

            body = cfg["body"]
            body_para = normalized.paragraphs[1]
            self.assertEqual(body_para.alignment, 3)  # WD_ALIGN_PARAGRAPH.JUSTIFY
            self.assertEqual(body_para.paragraph_format.line_spacing, body["lineSpacing"])
            self.assertEqual(body_para.paragraph_format.first_line_indent, Pt(24))

            table_cfg = cfg["table"]
            cell_para = normalized.tables[0].rows[0].cells[0].paragraphs[0]
            self.assertEqual(cell_para.paragraph_format.first_line_indent, Pt(0))
            self.assertEqual(cell_para.paragraph_format.line_spacing, table_cfg["cellLineSpacing"])

    def test_normalize_docx_rejects_invalid_paths_and_overwrite(self):
        from normalizer import normalize_docx

        cfg = load_config()
        source = ROOT / "规范文档示例.docx"

        with self.assertRaises(ValueError):
            normalize_docx(cfg, source, source)

        with tempfile.TemporaryDirectory() as tmp:
            bad_input = Path(tmp) / "not-word.txt"
            bad_input.write_text("plain text", encoding="utf-8")
            output = Path(tmp) / "out.docx"
            with self.assertRaises(ValueError):
                normalize_docx(cfg, bad_input, output)

    def test_normalize_docx_handles_no_heading_no_table_document(self):
        from normalizer import inspect_docx, normalize_docx

        cfg = load_config()
        with tempfile.TemporaryDirectory() as tmp:
            source = Path(tmp) / "plain.docx"
            output = Path(tmp) / "plain_规范化.docx"
            doc = Document()
            doc.add_paragraph("这是一份没有标题和表格的普通文档。")
            doc.add_paragraph("第二段正文用于检查段落数量。")
            doc.save(source)

            info = inspect_docx(source)
            normalize_docx(cfg, source, output)
            normalized = Document(output)

            self.assertEqual(info.table_count, 0)
            self.assertEqual(info.heading_levels, ())
            self.assertEqual(len(normalized.paragraphs), 2)
            self.assertEqual(normalized.paragraphs[0].paragraph_format.line_spacing, cfg["body"]["lineSpacing"])

    def test_normalize_docx_handles_multilevel_headings_and_margin_changes(self):
        from normalizer import inspect_docx, normalize_docx

        cfg = load_config()
        cfg["page"]["margins"]["left"] = 2.2
        with tempfile.TemporaryDirectory() as tmp:
            source = Path(tmp) / "headings.docx"
            output = Path(tmp) / "headings_规范化.docx"
            doc = Document()
            doc.add_paragraph("一、一级标题")
            doc.add_paragraph("（一）二级标题")
            doc.add_paragraph("1. 三级标题")
            doc.add_paragraph("（1）四级标题")
            doc.save(source)

            info = inspect_docx(source)
            normalize_docx(cfg, source, output)
            normalized = Document(output)

            self.assertEqual(info.heading_levels, (1, 2, 3, 4))
            self.assertAlmostEqual(normalized.sections[0].left_margin.cm, 2.2, places=1)
            self.assertEqual(normalized.paragraphs[0].style.name, "Heading 1")

    def test_normalize_docx_clears_first_line_chars_indent_in_table(self):
        """回归测试：单元格段落若有 OOXML 字符单位缩进（firstLineChars 等），
        规范化后必须清掉——否则 Word 仍按字符数渲染首行缩进。
        """
        from normalizer import normalize_docx

        cfg = load_config()
        with tempfile.TemporaryDirectory() as tmp:
            source = Path(tmp) / "table_with_firstlinechars.docx"
            output = Path(tmp) / "table_with_firstlinechars_规范化.docx"

            # 手工构造：单元格段落带 firstLine=480（24pt）+ firstLineChars=200（2字符）
            doc = Document()
            tbl = doc.add_table(rows=1, cols=1)
            tbl.style = "Table Grid"
            p = tbl.rows[0].cells[0].paragraphs[0]
            p.add_run("带字符单位缩进的单元格文字")
            pPr = p._element.get_or_add_pPr()
            ind = OxmlElement("w:ind")
            ind.set(qn("w:left"), "0")
            ind.set(qn("w:leftChars"), "0")
            ind.set(qn("w:firstLine"), "480")
            ind.set(qn("w:firstLineChars"), "200")
            ind.set(qn("w:right"), "0")
            pPr.append(ind)
            doc.save(source)

            normalize_docx(cfg, source, output)
            normalized = Document(output)
            cell_p = normalized.tables[0].rows[0].cells[0].paragraphs[0]
            pPr = cell_p._element.get_or_add_pPr()
            ind = pPr.find(qn("w:ind"))
            # 三个数值单位必须为 0
            self.assertEqual(ind.get(qn("w:firstLine")), "0")
            self.assertEqual(ind.get(qn("w:left")), "0")
            self.assertEqual(ind.get(qn("w:right")), "0")
            # 三个字符单位属性必须被清掉（OOXML 中字符单位优先于数值单位，残留会导致 bug）
            self.assertIsNone(ind.get(qn("w:firstLineChars")),
                              "firstLineChars 残留会导致 Word 仍按字符数缩进")
            self.assertIsNone(ind.get(qn("w:leftChars")),
                              "leftChars 残留会导致 Word 仍按字符数缩进")
            self.assertIsNone(ind.get(qn("w:rightChars")),
                              "rightChars 残留会导致 Word 仍按字符数缩进")

    def test_normalize_docx_clears_first_line_chars_in_normal_style(self):
        """回归测试：Normal 样式里的 firstLineChars 必须清掉。
        表格单元格段落没有显式 firstLineChars 时会继承 Normal 样式，
        若 Normal 样式 firstLineChars=200（2 字符），Word 仍按 2 字符渲染。
        """
        from normalizer import normalize_docx

        cfg = load_config()
        with tempfile.TemporaryDirectory() as tmp:
            source = Path(tmp) / "normal_with_firstlinechars.docx"
            output = Path(tmp) / "normal_with_firstlinechars_规范化.docx"

            # 手工构造：Normal 样式带 firstLineChars=200
            doc = Document()
            normal = doc.styles["Normal"]
            pPr = normal.element.find(qn("w:pPr"))
            if pPr is None:
                pPr = OxmlElement("w:pPr")
                normal.element.append(pPr)
            ind = OxmlElement("w:ind")
            ind.set(qn("w:firstLine"), "420")
            ind.set(qn("w:firstLineChars"), "200")
            pPr.append(ind)
            # 加一个表格 + 普通段落，模拟真实场景
            doc.add_paragraph("这是普通正文段落。")
            tbl = doc.add_table(rows=1, cols=1)
            tbl.rows[0].cells[0].paragraphs[0].add_run("单元格文字")
            doc.save(source)

            normalize_docx(cfg, source, output)
            normalized = Document(output)
            normal = normalized.styles["Normal"]
            pPr = normal.element.find(qn("w:pPr"))
            ind = pPr.find(qn("w:ind")) if pPr is not None else None
            # 关键断言：Normal 样式里 firstLineChars 必须清掉
            if ind is not None:
                self.assertIsNone(ind.get(qn("w:firstLineChars")),
                                  "Normal 样式的 firstLineChars 必须清掉，否则表格段落会继承")


def _add_page_field(paragraph):
    """在段落中插入 PAGE 域。"""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def _build_cover_toc_document(path):
    """构造一份含封面（无页码 section）+ 目录条目 + 正文的测试文档。

    结构：
      section 0（封面/前置页，无页码）：封面标题、文档信息、封面表格、分节符
      section 1（正文，有页码）：目录标题、TOC 条目（toc 1/2 样式）、一级标题、正文
    """
    doc = Document()
    try:
        doc.styles.add_style("toc 1", WD_STYLE_TYPE.PARAGRAPH)
        doc.styles.add_style("toc 2", WD_STYLE_TYPE.PARAGRAPH)
    except Exception:
        pass

    # 封面 section（section 0）
    doc.add_paragraph("封面标题")
    doc.add_paragraph("文档信息")
    cover_tbl = doc.add_table(rows=2, cols=2)
    cover_tbl.rows[0].cells[0].text = "项目"
    cover_tbl.rows[0].cells[1].text = "编号"
    cover_tbl.rows[1].cells[0].text = "名称"
    cover_tbl.rows[1].cells[1].text = "示例"
    doc.add_paragraph()  # 承载分节符的空段落
    doc.add_section(WD_SECTION.NEW_PAGE)

    # 正文 section（section 1）
    doc.add_paragraph("目录")  # 目录标题
    p_toc1 = doc.add_paragraph("1. 引言\t1")
    try:
        p_toc1.style = "toc 1"
    except Exception:
        pass
    p_toc2 = doc.add_paragraph("1.1 目的\t1")
    try:
        p_toc2.style = "toc 2"
    except Exception:
        pass
    doc.add_paragraph("一、引言")  # 一级标题
    doc.add_paragraph("正文内容。")  # 正文

    # 仅给正文 section 加页码，封面 section 不加
    sec1 = doc.sections[1]
    sec1.footer.is_linked_to_previous = False
    footer_p = sec1.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_page_field(footer_p)

    doc.save(path)
    return doc


class NormalizeCoverAndTocTests(unittest.TestCase):
    """验证封面保留与目录错乱修复。"""

    def test_normalize_preserves_cover_and_toc(self):
        from normalizer import normalize_docx

        cfg = load_config()
        with tempfile.TemporaryDirectory() as tmp:
            source = Path(tmp) / "with_cover.docx"
            output = Path(tmp) / "with_cover_规范化.docx"
            _build_cover_toc_document(source)

            normalize_docx(cfg, source, output)
            result = Document(output)

            # 段落索引（add_section 会插入一个空段作为分节符载体）：
            # 0 封面标题 / 1 文档信息 / 2 空段 / 3 分节符空段 / 4 目录 / 5 toc1 / 6 toc2 / 7 H1 / 8 正文
            # 1) 目录标题保留原格式，未被套正文格式（正文会加 Pt(24) 首行缩进）
            title_para = result.paragraphs[4]
            self.assertEqual((title_para.text or "").strip(), "目录")
            self.assertNotEqual(title_para.paragraph_format.first_line_indent, Pt(24))

            # 2) TOC 条目样式保留，未被改成 Heading
            toc1 = result.paragraphs[5]
            toc2 = result.paragraphs[6]
            self.assertTrue(toc1.style.name.lower().startswith("toc"))
            self.assertTrue(toc2.style.name.lower().startswith("toc"))

            # 3) 封面段落保留原格式，未被套正文格式
            cover_para = result.paragraphs[1]
            self.assertEqual((cover_para.text or "").strip(), "文档信息")
            self.assertNotEqual(cover_para.paragraph_format.first_line_indent, Pt(24))

            # 4) 封面 section 无页码，正文 section 有页码
            self.assertNotIn("PAGE", result.sections[0].footer._element.xml)
            self.assertIn("PAGE", result.sections[1].footer._element.xml)

            # 5) 正文标题与正文段落正常规范化
            h1_para = result.paragraphs[7]
            self.assertEqual(h1_para.style.name, "Heading 1")
            body_para = result.paragraphs[8]
            self.assertEqual(body_para.alignment, 3)  # JUSTIFY
            self.assertEqual(body_para.paragraph_format.first_line_indent, Pt(24))

    def test_normalize_preserves_cover_table(self):
        from normalizer import normalize_docx

        cfg = load_config()
        with tempfile.TemporaryDirectory() as tmp:
            source = Path(tmp) / "with_cover.docx"
            output = Path(tmp) / "with_cover_规范化.docx"
            _build_cover_toc_document(source)

            normalize_docx(cfg, source, output)
            result = Document(output)

            # 封面表格保留原版式：单元格段落对齐未被改为 center（正文表格默认会居中）
            cover_cell = result.tables[0].rows[0].cells[0]
            self.assertNotEqual(cover_cell.paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.CENTER)

    def test_mark_toc_fields_dirty(self):
        from normalizer import _mark_toc_fields_dirty

        doc = Document()
        p = doc.add_paragraph()
        r1 = p.add_run()
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        r1._r.append(begin)
        r2 = p.add_run()
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = 'TOC \\o "1-3" \\h \\u '
        r2._r.append(instr)
        r3 = p.add_run()
        sep = OxmlElement("w:fldChar")
        sep.set(qn("w:fldCharType"), "separate")
        r3._r.append(sep)
        r4 = p.add_run()
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        r4._r.append(end)

        _mark_toc_fields_dirty(doc)
        self.assertEqual(begin.get(qn("w:dirty")), "true")

    def test_helpers_section_and_toc_detection(self):
        from normalizer import (
            _build_body_section_maps,
            _collect_cover_section_indices,
            _find_toc_title_index,
            _is_toc_style,
            _section_has_page_number,
        )

        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "with_cover.docx"
            _build_cover_toc_document(path)
            doc = Document(path)

            # section 0 无页码（封面），section 1 有页码（正文）
            self.assertFalse(_section_has_page_number(doc.sections[0]))
            self.assertTrue(_section_has_page_number(doc.sections[1]))

            cover = _collect_cover_section_indices(doc)
            self.assertEqual(cover, {0})

            para_to_section, table_to_section = _build_body_section_maps(doc)
            # 封面段落属于 section 0，正文段落属于 section 1
            self.assertEqual(para_to_section[0], 0)
            self.assertEqual(para_to_section[4], 1)
            # 封面表格属于 section 0
            self.assertEqual(table_to_section[0], 0)

            # 目录标题索引应为 4（首个 toc 条目 5 之前最近的非空非 toc 段）
            self.assertEqual(_find_toc_title_index(doc), 4)

            # toc 样式识别
            self.assertTrue(_is_toc_style(doc.paragraphs[5]))
            self.assertFalse(_is_toc_style(doc.paragraphs[7]))


class NormalizeErrorTests(unittest.TestCase):
    def test_normalize_error_includes_message_cause_hint(self):
        from normalizer import NormalizeError
        e = NormalizeError("找不到文件", cause=FileNotFoundError("x"), hint="检查路径")
        s = str(e)
        self.assertIn("找不到文件", s)
        self.assertIn("FileNotFoundError", s)
        self.assertIn("检查路径", s)

    def test_input_not_found_is_normalize_error(self):
        from normalizer import NormalizeError, InputNotFoundError
        self.assertTrue(issubclass(InputNotFoundError, NormalizeError))

    def test_each_subclass_exists(self):
        from normalizer import (
            NormalizeError, InputNotFoundError, InvalidFileTypeError,
            CorruptDocxError, OutputNotWritableError, SameInputOutputError,
        )
        for cls in (InputNotFoundError, InvalidFileTypeError, CorruptDocxError,
                    OutputNotWritableError, SameInputOutputError):
            self.assertTrue(issubclass(cls, NormalizeError))


if __name__ == "__main__":
    unittest.main()
