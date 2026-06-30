"""
formatter.py —— 规范文档排版核心库（Python / python-docx，配置驱动）
=================================================================
依据 config.json 生成符合规范的 .docx。所有字体、字号、行距、边距、对齐
均由配置决定。

关键修复：表格文字完全独立于正文——
  1) Normal 默认样式不设首行缩进、用单倍行距（正文 1.5 倍与首行缩进改为
     只在正文段落上显式设置，不落在 Normal 上）；
  2) 表格单元格段落显式置 first_line_indent=0、line_spacing=1.0、
     space_before/after=0，并使用独立的单元格字体/字号，
     杜绝“表格文字引用正文文本导致首行缩进和段间距”的问题。

字号写法：中文名（小二/小三/小四/五号/小五…）或磅值数字（18/15/12/10.5/9）。
"""

import os
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ------------------------------------------------------------------
# 常量与解析
# ------------------------------------------------------------------

CM_PER_INCH = 2.54
PT_PER_CM = 28.3464567  # 1cm = 28.35 pt
EMU_PER_CM = 360000

SIZE_NAME_TO_PT = {
    "初号": 42, "小初": 36,
    "一号": 26, "小一": 24,
    "二号": 22, "小二": 18,
    "三号": 16, "小三": 15,
    "四号": 14, "小四": 12,
    "五号": 10.5, "小五": 9,
    "六号": 7.5, "小六": 6.5,
    "七号": 5.5, "八号": 5,
}

ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    "both": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def resolve_size(val):
    """字号名/磅值 → pt。默认小四(12)。"""
    if val is None:
        return 12.0
    if isinstance(val, (int, float)):
        return float(val)
    s = str(val).strip()
    if s in SIZE_NAME_TO_PT:
        return SIZE_NAME_TO_PT[s]
    try:
        return float(s)
    except ValueError:
        return 12.0


def resolve_align(a):
    return ALIGN_MAP.get(str(a).lower() if a else "left", WD_ALIGN_PARAGRAPH.LEFT)


def set_run_font(run, zh_font, en_font, size_pt=None, bold=None, color=None):
    """设置 run 字体：中文 eastAsia + 英数 ascii，并处理字号/加粗/颜色。"""
    en_font = en_font or "Times New Roman"
    run.font.name = en_font
    # 确保 rPr/rFonts 存在并设置 eastAsia
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), en_font)
    rfonts.set(qn("w:hAnsi"), en_font)
    rfonts.set(qn("w:cs"), en_font)
    rfonts.set(qn("w:eastAsia"), zh_font or "宋体")
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_para_spacing(para, line_spacing_mult=None, before_pt=None, after_pt=None):
    """设置段落行距与段前段后。line_spacing_mult 为倍数（1.0/1.5）。"""
    pf = para.paragraph_format
    if line_spacing_mult is not None:
        pf.line_spacing = line_spacing_mult
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    if before_pt is not None:
        pf.space_before = Pt(before_pt)
    if after_pt is not None:
        pf.space_after = Pt(after_pt)


def set_zero_indent(para):
    """显式清零所有缩进，避免继承正文的首行缩进。

    同时清掉 OOXML 字符单位属性（firstLineChars/leftChars/rightChars）：
    这些属性在 OOXML 中优先于数值单位（firstLine/left/right），仅清数值
    单位会留下"2 字符缩进"等残留，导致 Word 仍按字符数渲染。
    """
    pf = para.paragraph_format
    pf.first_line_indent = Pt(0)
    pf.left_indent = Pt(0)
    pf.right_indent = Pt(0)
    pPr = para._element.get_or_add_pPr()
    ind = pPr.find(qn("w:ind"))
    if ind is not None:
        for attr in ("firstLineChars", "leftChars", "rightChars"):
            key = qn("w:" + attr)
            if key in ind.attrib:
                del ind.attrib[key]


def clear_style_indent_chars(style):
    """清掉样式（通常是 Normal）<w:ind> 里的字符单位属性。

    set_zero_indent 处理的是段落级别的 <w:ind>，但段落会从样式继承 firstLineChars。
    若样式（如 Normal）里 firstLineChars=200 而段落自己没显式声明，Word 仍按
    2 字符渲染首行缩进。此函数专用于清掉样式里的字符单位属性。
    """
    pPr = style.element.find(qn("w:pPr"))
    if pPr is None:
        return
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        return
    for attr in ("firstLineChars", "leftChars", "rightChars"):
        key = qn("w:" + attr)
        if key in ind.attrib:
            del ind.attrib[key]


def add_page_number_field(paragraph, run_font_fn):
    """在段落里插入 PAGE 域（页码）。"""
    run = paragraph.add_run()
    run_font_fn(run)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def get_image_size(path):
    """读取图片像素尺寸（用 python-docx 内置 image 模块）。"""
    try:
        from docx.image.image import Image
        img = Image.from_file(path)
        return img.px_width, img.px_height
    except Exception:
        return None, None


# ------------------------------------------------------------------
# Formatter —— 核心类
# ------------------------------------------------------------------

class Formatter:
    def __init__(self, config):
        self.cfg = config
        self.doc = Document()
        self.fig_no = 0
        self.tab_no = 0
        self._setup_page()
        self._setup_styles()

    # ---- 页面与样式 ----
    def _setup_page(self):
        cfg = self.cfg
        sec = self.doc.sections[0]
        # A4
        sec.page_width = Cm(21.0)
        sec.page_height = Cm(29.7)
        m = cfg["page"]["margins"]
        sec.top_margin = Cm(m["top"])
        sec.bottom_margin = Cm(m["bottom"])
        sec.left_margin = Cm(m["left"])
        sec.right_margin = Cm(m["right"])
        self.content_width_cm = 21.0 - m["left"] - m["right"]

        # 页脚页码（底部居中）
        pn = cfg.get("pageNumber", {})
        if pn.get("enabled", True):
            footer = sec.footer
            footer.is_linked_to_previous = False
            p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            p.alignment = resolve_align(pn.get("align", "center"))
            set_para_spacing(p, line_spacing_mult=1.0, before_pt=0, after_pt=0)
            set_zero_indent(p)
            add_page_number_field(p, lambda r: set_run_font(
                r, pn.get("font", "宋体"), pn.get("asciiFont", "Times New Roman"),
                size_pt=resolve_size(pn.get("size", "小五"))))

    def _setup_styles(self):
        """Normal 样式：单倍行距、无首行缩进（正文属性只在正文段落显式设置）。"""
        cfg = self.cfg
        normal = self.doc.styles["Normal"]
        body = cfg["body"]
        set_run_font_element(normal.element, body["font"], body["asciiFont"],
                             resolve_size(body["size"]), body.get("color", "000000"))
        normal.font.size = Pt(resolve_size(body["size"]))
        npf = normal.paragraph_format
        npf.line_spacing = 1.0  # Normal 用单倍，避免表格/页码继承 1.5 倍
        npf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        npf.first_line_indent = Pt(0)
        npf.space_before = Pt(0)
        npf.space_after = Pt(0)
        # 关键：清掉 Normal 样式 <w:ind> 里的字符单位属性（firstLineChars 等），
        # 否则表格单元格段落继承后会按字符数渲染首行缩进。
        clear_style_indent_chars(normal)

        # 标题样式覆盖（黑体/宋体 + 黑色 + 对应字号，覆盖内置蓝色默认）
        for lv in range(1, 5):
            key = "h" + str(lv)
            hc = cfg["headings"].get(key)
            if not hc:
                continue
            try:
                hs = self.doc.styles["Heading " + str(lv)]
            except KeyError:
                continue
            set_run_font_element(hs.element, hc["font"], hc["asciiFont"],
                                 resolve_size(hc["size"]), hc.get("color", "000000"))
            hs.font.size = Pt(resolve_size(hc["size"]))
            hs.font.bold = hc.get("bold", True)
            hs.font.color.rgb = RGBColor.from_string(hc.get("color", "000000"))
            hpf = hs.paragraph_format
            hpf.line_spacing = hc.get("lineSpacing", 1.5)
            hpf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            hpf.space_before = Pt(hc.get("spaceBeforePt", 0))
            hpf.space_after = Pt(hc.get("spaceAfterPt", 0))

    # ---- 正文 ----
    def body_para(self, text):
        cfg = self.cfg["body"]
        p = self.doc.add_paragraph()
        p.alignment = resolve_align(cfg.get("align", "justify"))
        # 正文：显式 1.5 倍行距 + 首行缩进（不靠 Normal 继承）
        set_para_spacing(p, line_spacing_mult=cfg.get("lineSpacing", 1.5),
                         before_pt=0, after_pt=cfg.get("spaceAfterPt", 0))
        indent_chars = cfg.get("firstLineIndentChars", 0)
        if indent_chars > 0:
            p.paragraph_format.first_line_indent = Pt(resolve_size(cfg["size"]) * indent_chars)
        run = p.add_run(text)
        set_run_font(run, cfg["font"], cfg["asciiFont"],
                     size_pt=resolve_size(cfg["size"]), color=cfg.get("color", "000000"))
        return p

    # ---- 标题 ----
    def heading(self, level, text):
        hc = self.cfg["headings"]["h" + str(level)]
        p = self.doc.add_paragraph(style="Heading " + str(level))
        p.alignment = resolve_align(hc.get("align", "left"))
        set_para_spacing(p, line_spacing_mult=hc.get("lineSpacing", 1.5),
                         before_pt=hc.get("spaceBeforePt", 0), after_pt=hc.get("spaceAfterPt", 0))
        set_zero_indent(p)
        run = p.add_run(text)
        set_run_font(run, hc["font"], hc["asciiFont"],
                     size_pt=resolve_size(hc["size"]),
                     bold=hc.get("bold", True), color=hc.get("color", "000000"))
        return p

    def h1(self, t): return self.heading(1, t)
    def h2(self, t): return self.heading(2, t)
    def h3(self, t): return self.heading(3, t)
    def h4(self, t): return self.heading(4, t)

    # ---- 图 ----
    def figure(self, image_path, title, source=None):
        f = self.cfg["figure"]
        self.fig_no += 1
        no = self.fig_no
        cap_font = f.get("captionText", "宋体")
        cap_en = f.get("captionAsciiFont", self.cfg["body"]["asciiFont"])
        out = []

        # 图片尺寸：默认宽度=正文宽度，按宽高比缩放，超高则反向缩放
        w_px, h_px = get_image_size(image_path)
        width_cm = f.get("maxWidthCm", 15.5)
        height_cm = None
        if w_px and h_px:
            ratio = w_px / h_px
            height_cm = width_cm / ratio
            max_h = f.get("maxHeightCm", 22)
            if height_cm > max_h:
                height_cm = max_h
                width_cm = height_cm * ratio

        # 图题段
        def cap_para():
            p = self.doc.add_paragraph()
            p.alignment = resolve_align(f.get("captionAlign", "center"))
            # 图题单倍行距 + 零缩进（修复行距导致居中漂移）
            set_para_spacing(p, line_spacing_mult=f.get("captionLineSpacing", 1.0),
                             before_pt=0, after_pt=2)
            set_zero_indent(p)
            run = p.add_run("%s%d  %s" % (f.get("prefix", "图"), no, title))
            set_run_font(run, cap_font, cap_en, size_pt=resolve_size(f.get("captionSize", "五号")),
                         bold=f.get("captionBold", True))
            return p

        # 来源/注释段
        def note_para(src):
            p = self.doc.add_paragraph()
            p.alignment = resolve_align(f.get("noteAlign", "left"))
            set_para_spacing(p, line_spacing_mult=f.get("noteLineSpacing", 1.0),
                             before_pt=0, after_pt=9)
            set_zero_indent(p)
            run = p.add_run(src)
            set_run_font(run, cap_font, cap_en, size_pt=resolve_size(f.get("noteSize", "小五")))
            return p

        # 图片段
        def img_para():
            p = self.doc.add_paragraph()
            p.alignment = resolve_align(f.get("align", "center"))
            set_para_spacing(p, line_spacing_mult=1.0, before_pt=6, after_pt=3)
            set_zero_indent(p)
            run = p.add_run()
            kwargs = {"width": Cm(width_cm)}
            if height_cm:
                kwargs["height"] = Cm(height_cm)
            run.add_picture(image_path, **kwargs)
            return p

        # 顺序：above → 图题→图片→注释；below → 图片→图题→注释
        if f.get("captionPosition", "below") == "above":
            out.append(cap_para())
            out.append(img_para())
            if source:
                out.append(note_para(source))
        else:
            out.append(img_para())
            out.append(cap_para())
            if source:
                out.append(note_para(source))
        return out

    # ---- 表 ----
    def table(self, title, headers, rows, note=None, column_widths_cm=None):
        t = self.cfg["table"]
        self.tab_no += 1
        no = self.tab_no
        cap_font = t.get("captionText", "宋体")
        cap_en = t.get("captionAsciiFont", self.cfg["body"]["asciiFont"])
        cell_font = t.get("cellFont", self.cfg["body"]["font"])
        cell_en = t.get("cellAsciiFont", self.cfg["body"]["asciiFont"])

        # 表题段（五号加粗居中，单倍行距 + 零缩进）
        def make_caption():
            p = self.doc.add_paragraph()
            p.alignment = resolve_align(t.get("captionAlign", "center"))
            set_para_spacing(p, line_spacing_mult=t.get("captionLineSpacing", 1.0),
                             before_pt=9, after_pt=4)
            set_zero_indent(p)
            run = p.add_run("%s%d  %s" % (t.get("prefix", "表"), no, title))
            set_run_font(run, cap_font, cap_en, size_pt=resolve_size(t.get("captionSize", "五号")),
                         bold=t.get("captionBold", True))
            return p

        # 注释/来源段（小五左对齐）
        def make_note():
            p = self.doc.add_paragraph()
            p.alignment = resolve_align(t.get("noteAlign", "left"))
            set_para_spacing(p, line_spacing_mult=t.get("noteLineSpacing", 1.0),
                             before_pt=3, after_pt=9)
            set_zero_indent(p)
            run = p.add_run(note)
            set_run_font(run, cap_font, cap_en, size_pt=resolve_size(t.get("noteSize", "小五")))
            return p

        # 顺序：above(默认) → 表题→表→注释；below → 表→表题→注释
        # python-docx 按创建顺序追加，故按目标顺序创建元素（无需事后移动）
        caption_position = t.get("captionPosition", "above")
        if caption_position == "above":
            make_caption()

        # 表格
        ncol = len(headers)
        tbl = self.doc.add_table(rows=1 + len(rows), cols=ncol)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.style = "Table Grid"
        # 列宽
        total_cm = self.content_width_cm
        if column_widths_cm and len(column_widths_cm) == ncol:
            widths = column_widths_cm
        else:
            widths = [total_cm / ncol] * ncol
            widths[-1] = total_cm - sum(widths[:-1])
        for i, w in enumerate(widths):
            for row in tbl.rows:
                row.cells[i].width = Cm(w)

        cell_align = resolve_align(t.get("cellAlign", "center"))
        cell_valign = WD_ALIGN_VERTICAL.CENTER if t.get("cellVAlign", "center") == "center" else WD_ALIGN_VERTICAL.TOP
        cell_size = resolve_size(t.get("cellSize", "小四"))
        header_bold = t.get("headerBold", True)

        def fill_cell(cell, text, bold):
            cell.vertical_alignment = cell_valign
            p = cell.paragraphs[0]
            p.clear()
            p.alignment = cell_align
            # 【关键修复】单元格段落完全独立：单倍行距 + 零缩进 + 零段距
            set_para_spacing(p, line_spacing_mult=t.get("cellLineSpacing", 1.0),
                             before_pt=0, after_pt=0)
            set_zero_indent(p)
            run = p.add_run(str(text))
            set_run_font(run, cell_font, cell_en, size_pt=cell_size, bold=bold)

        # 表头
        for i, h in enumerate(headers):
            fill_cell(tbl.rows[0].cells[i], h, header_bold)
        # 数据
        for r, row in enumerate(rows):
            for i, val in enumerate(row):
                fill_cell(tbl.rows[1 + r].cells[i], val, False)

        # 表题在下时，表后创建表题
        if caption_position == "below":
            make_caption()

        # 注释/来源
        if note:
            make_note()
        return tbl

    # ---- 参考文献 GB/T 7714 ----
    def references(self, items):
        r = self.cfg["references"]
        hang_pt = r.get("hangingIndentPt", 24)
        size_pt = resolve_size(r.get("size", "小五"))
        for i, it in enumerate(items):
            n = i + 1
            s = "[%d] %s. %s[%s]. " % (n, it["authors"], it["title"], it["type"])
            if it.get("place") and it.get("publisher"):
                s += "%s: %s, %s." % (it["place"], it["publisher"], it.get("year", ""))
            elif it.get("source"):
                s += "%s, %s" % (it["source"], it.get("year", ""))
                if it.get("volume"):
                    s += ", %s" % it["volume"]
                if it.get("issue"):
                    s += "(%s)" % it["issue"]
                s += ": %s." % it.get("pages", "")
            else:
                s += "%s." % it.get("year", "")
            if it.get("url"):
                s += " %s" % it["url"]
                if it.get("accessDate"):
                    s += " [%s]." % it["accessDate"]
            p = self.doc.add_paragraph()
            p.alignment = resolve_align(r.get("align", "justify"))
            set_para_spacing(p, line_spacing_mult=r.get("lineSpacing", 1.0), before_pt=0, after_pt=3)
            pf = p.paragraph_format
            pf.left_indent = Pt(hang_pt)
            pf.first_line_indent = Pt(-hang_pt)
            run = p.add_run(s)
            set_run_font(run, r.get("font", "宋体"), r.get("asciiFont", "Times New Roman"),
                         size_pt=size_pt)
        return

    def save(self, path):
        self.doc.save(path)


def set_run_font_element(style_element, zh_font, en_font, size_pt, color):
    """在样式元素上设置 rFonts（含 eastAsia）、字号、颜色。"""
    rpr = style_element.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        style_element.append(rpr)
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), en_font or "Times New Roman")
    rfonts.set(qn("w:hAnsi"), en_font or "Times New Roman")
    rfonts.set(qn("w:cs"), en_font or "Times New Roman")
    rfonts.set(qn("w:eastAsia"), zh_font or "宋体")
    # 清掉可能残留的 color（避免内置蓝色）
    for c in rpr.findall(qn("w:color")):
        rpr.remove(c)
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), color or "000000")
    rpr.append(color_el)


# ------------------------------------------------------------------
# 示例内容 + 生成入口
# ------------------------------------------------------------------

def build_sample_content(F, assets_dir):
    F.h1("一、研究背景与意义")
    F.body_para(
        "随着信息技术的快速发展，规范化学术文档的排版质量日益受到重视。本文档用于演示一套可复用、可配置的排版规范，"
        "覆盖正文、标题、图、表与参考文献等核心要素，确保生成结果在 Word 中保持一致的视觉表现。"
    )
    F.body_para(
        "其中，图表标题的居中问题与表格文字继承正文格式的问题是历史方案的主要痛点。本方案将图题/表题改用单倍行距并"
        "显式置零缩进，同时令表格单元格段落完全独立于正文（独立字体、单倍行距、零缩进、零段距），从根源上消除首行"
        "缩进与段间距的串扰。"
    )

    F.h2("（一）图的排版规范")
    F.body_para(
        "图片统一宽度等于正文宽度，高度按原始宽高比等比缩放，超高时以高度上限反向缩放；图片水平居中、不加边框，"
        "图题置于图下方，五号加粗居中，来源或注释以小五号左对齐排列。"
    )
    F.figure(os.path.join(assets_dir, "figure1.png"), "各季度销售额对比",
             source="数据来源：示例数据，仅用于排版演示。")

    F.h2("（二）表的排版规范")
    F.body_para(
        "表题置于表上方，五号加粗居中；表格整体居中并占满正文宽度，单元格内容垂直且水平居中；注释或来源置于表下方，"
        "小五号左对齐。表格文字使用独立格式，不继承正文的首行缩进与 1.5 倍行距。"
    )
    F.table("主要指标年度对比",
            ["指标", "2023 年", "2024 年", "同比增长"],
            [["营业收入（万元）", "1,200", "1,500", "25.0%"],
             ["净利润（万元）", "180", "225", "25.0%"],
             ["研发投入（万元）", "96", "135", "40.6%"]],
            note="注：以上数据为示例，不构成任何投资建议。")

    F.h2("（三）三级标题演示")
    F.h3("1. 三级标题示例")
    F.body_para("三级标题采用宋体小三，用于更细的章节划分。")

    F.h1("二、参考文献著录规范")
    F.body_para(
        "参考文献遵循 GB/T 7714 顺序编码制：正文引用以 [n] 编号，逾 3 位作者用“等”或“et al”；强制标注文献类型标识，"
        "电子文献使用复合标识（如 [J/OL]），网络资源须含引用日期。"
    )
    F.references([
        {"authors": "王建国, 李明, 张华, 等", "title": "现代排版工程", "type": "M",
         "place": "北京", "publisher": "科学出版社", "year": "2022"},
        {"authors": "LIU Y, CHEN X, ZHAO M, et al", "title": "A study on document formatting consistency",
         "type": "J", "source": "Journal of Publishing Science", "year": "2023",
         "volume": "15", "issue": "3", "pages": "45-52"},
        {"authors": "陈晓", "title": "学术写作中的图表规范研究", "type": "D",
         "place": "上海", "publisher": "复旦大学", "year": "2021"},
        {"authors": "全国信息与文献标准化技术委员会",
         "title": "信息与文献 参考文献著录规则: GB/T 7714—2015", "type": "S",
         "place": "北京", "publisher": "中国标准出版社", "year": "2015"},
        {"authors": "张伟, 刘洋", "title": "电子文献引用规范的新进展", "type": "J/OL",
         "source": "中国科技期刊研究", "year": "2024", "volume": "35", "issue": "2",
         "pages": "88-95", "url": "https://doi.org/10.1234/example.2024.02.012",
         "accessDate": "2026-06-28"},
        {"authors": "中华人民共和国国家统计局",
         "title": "2024 年国民经济和社会发展统计公报", "type": "EB/OL",
         "url": "https://www.stats.gov.cn/example.htm", "accessDate": "2026-06-28"},
    ])


def generate(config, output_path, assets_dir=None, content_fn=build_sample_content):
    """根据 config 生成 docx 到 output_path。"""
    if assets_dir is None:
        assets_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "assets")
    F = Formatter(config)
    content_fn(F, assets_dir)
    F.save(output_path)
    return F


if __name__ == "__main__":
    import json
    cfg_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "config.json")
    with open(cfg_path, "r", encoding="utf-8") as f:
        config = json.load(f)
    out = config.get("output", {}).get("filename", "规范文档示例.docx")
    out_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), out)
    generate(config, out_path)
    print("已生成: " + out_path)
    print("正文宽度 ≈ %.2f cm" % (21.0 - config["page"]["margins"]["left"] - config["page"]["margins"]["right"]))
