"""Conservative normalization for existing Word documents."""

from dataclasses import dataclass, field
import os
import re
import zipfile

import lxml.etree

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn

from formatter import (
    resolve_size,
    resolve_align,
    set_run_font,
    set_para_spacing,
    set_zero_indent,
    clear_style_indent_chars,
    add_page_number_field,
    set_run_font_element,
)


class NormalizeError(Exception):
    """规范化失败的统一异常，message 直接给最终用户看。"""

    def __init__(self, message, *, cause=None, hint=None):
        super().__init__(message)
        self.message = message
        self.cause = cause
        self.hint = hint

    def __str__(self):
        parts = [self.message]
        if self.cause is not None:
            parts.append(f"原因：{type(self.cause).__name__}: {self.cause}")
        if self.hint:
            parts.append(f"建议：{self.hint}")
        return "\n".join(parts)


class InputNotFoundError(NormalizeError):
    """输入文件不存在。"""


class InvalidFileTypeError(NormalizeError):
    """输入不是 .docx 文件。"""


class CorruptDocxError(NormalizeError):
    """docx 文件损坏（zip 或 XML 解析失败）。"""


class OutputNotWritableError(NormalizeError):
    """输出路径不可写（权限/磁盘）。"""


class SameInputOutputError(NormalizeError):
    """输出路径与输入相同。"""


class WarningCollector:
    """规范化过程中收集非致命警告。"""

    def __init__(self):
        self.warnings = []

    def warn(self, msg, *, location=None):
        prefix = f"[{location}] " if location else ""
        self.warnings.append(f"{prefix}{msg}")

    def __enter__(self):
        return self

    def __exit__(self, *args):
        return None


@dataclass
class NormalizeResult:
    """规范化结果。dry_run=True 时 output_path 为 None，backup_path 为 None。"""
    input_path: str
    output_path: "str | None" = None
    backup_path: "str | None" = None
    dry_run: bool = False
    paragraphs_processed: int = 0
    tables_processed: int = 0
    images_processed: int = 0
    changes: list = field(default_factory=list)
    warnings: list = field(default_factory=list)
    errors: list = field(default_factory=list)
    duration_ms: int = 0


EMU_PER_CM = 360000

HEADING_PATTERNS = {
    1: re.compile(r"^[一二三四五六七八九十百]+、"),
    2: re.compile(r"^（[一二三四五六七八九十百]+）|^\([一二三四五六七八九十百]+\)"),
    3: re.compile(r"^\d+[.、]"),
    4: re.compile(r"^（\d+）|^\(\d+\)"),
}


@dataclass(frozen=True)
class DocxInspection:
    paragraph_count: int
    table_count: int
    image_count: int
    heading_levels: tuple
    has_header_footer: bool
    has_complex_objects: bool


def inspect_docx(input_path):
    if not _is_docx_path(input_path):
        raise ValueError("输入文件必须是 .docx Word 文档")
    doc = Document(input_path)
    heading_levels = sorted({_detect_heading_level(p) for p in doc.paragraphs} - {0})
    return DocxInspection(
        paragraph_count=len(doc.paragraphs),
        table_count=len(doc.tables),
        image_count=len(doc.inline_shapes),
        heading_levels=tuple(heading_levels),
        has_header_footer=_has_header_footer(doc),
        has_complex_objects=_has_complex_objects(doc),
    )


def normalize_docx(config, input_path, output_path, assets_dir=None, *,
                   backup=True, dry_run=False, return_result=False, on_warning=None):
    """规范化 docx。薄包装层。

    旧签名（不传新参数）行为完全兼容：返回 str(output_path)。
    新签名支持 backup / dry_run / return_result / on_warning。
    抛 NormalizeError 子类替代原本的 ValueError / FileNotFoundError / KeyError。
    """
    import time
    del assets_dir  # Kept for a stable public interface.
    start = time.time()
    input_path = os.fspath(input_path)
    output_path = os.fspath(output_path)

    # 1. 路径校验（抛 NormalizeError 子类）
    _validate_paths(input_path, output_path)

    # 2. 备份（dry_run 时跳过；备份失败归为 warning）
    backup_path = None
    if backup and not dry_run:
        backup_path = _backup_source(input_path)
        if backup_path is None:
            warn_msg = "备份失败（不影响规范化继续）"
            if on_warning:
                on_warning(warn_msg, location="备份")

    # 3. 加载 + 规范化（try/except 把异常转成 NormalizeError）
    with WarningCollector() as wc:
        try:
            doc = Document(input_path)
        except (zipfile.BadZipFile, lxml.etree.XMLSyntaxError) as e:
            raise CorruptDocxError(
                f"docx 文件已损坏: {input_path}",
                cause=e,
                hint="用 Word 打开重新保存一次",
            ) from e
        except KeyError as e:
            raise CorruptDocxError(
                f"docx 关键样式缺失: {input_path}",
                cause=e,
                hint="用 Word 打开重新保存一次",
            ) from e

        try:
            para_to_section, table_to_section = _build_body_section_maps(doc)
            cover_sections = _collect_cover_section_indices(doc)

            _apply_page_settings(doc, config)
            _apply_document_styles(doc, config)
            _normalize_paragraphs(doc, config, para_to_section, cover_sections, on_warning=wc.warn)
            _normalize_tables(doc, config, table_to_section, cover_sections, on_warning=wc.warn)
            _normalize_inline_images(doc, config, on_warning=wc.warn)
            _mark_toc_fields_dirty(doc)
        except PermissionError as e:
            raise OutputNotWritableError(
                f"无法访问文件: {output_path}",
                cause=e,
                hint="检查目录权限和磁盘空间",
            ) from e

        # 4. 写文件（dry_run 时跳过）
        if not dry_run:
            try:
                out_dir = os.path.dirname(os.path.abspath(output_path))
                if out_dir:
                    os.makedirs(out_dir, exist_ok=True)
                doc.save(output_path)
            except PermissionError as e:
                raise OutputNotWritableError(
                    f"无法写入输出: {output_path}",
                    cause=e,
                    hint="检查目录权限和磁盘空间",
                ) from e
            except OSError as e:
                raise OutputNotWritableError(
                    f"无法写入输出: {output_path}",
                    cause=e,
                    hint="检查目录权限和磁盘空间",
                ) from e

    duration_ms = int((time.time() - start) * 1000)
    if return_result:
        return NormalizeResult(
            input_path=input_path,
            output_path=output_path if not dry_run else None,
            backup_path=backup_path,
            dry_run=dry_run,
            paragraphs_processed=len(doc.paragraphs),
            tables_processed=len(doc.tables),
            images_processed=len(getattr(doc, "inline_shapes", [])),
            changes=[],
            warnings=list(wc.warnings),
            errors=[],
            duration_ms=duration_ms,
        )
    return output_path


def _same_path(a, b):
    try:
        return os.path.abspath(os.fspath(a)).lower() == os.path.abspath(os.fspath(b)).lower()
    except TypeError:
        return False


def _is_docx_path(path):
    return str(path).lower().endswith(".docx")


def _validate_paths(input_path, output_path):
    """校验输入输出路径，抛 NormalizeError 子类。"""
    input_path = os.fspath(input_path)
    output_path = os.fspath(output_path)

    if not _is_docx_path(input_path):
        raise InvalidFileTypeError(
            f"不是有效的 docx 文件: {input_path}",
            hint="确认文件是 Word 导出的 .docx 格式（不是 .doc / .rtf / .wps）",
        )
    if not _is_docx_path(output_path):
        raise InvalidFileTypeError(
            f"输出路径不是 .docx 后缀: {output_path}",
            hint="输出文件必须以 .docx 结尾",
        )
    if not os.path.exists(input_path):
        raise InputNotFoundError(
            f"找不到文件: {input_path}",
            hint="检查路径是否正确，文件名是否包含特殊字符",
        )
    if _same_path(input_path, output_path):
        raise SameInputOutputError(
            f"输出与输入相同: {input_path}",
            hint="选择不同的输出文件",
        )


def _backup_source(input_path):
    """把输入文件复制到同目录的 .bak（不覆盖已有 .bak）。失败返回 None。"""
    import shutil
    input_path = os.fspath(input_path)
    if not os.path.exists(input_path):
        return None
    try:
        base = input_path + ".bak"
        target = base
        n = 0
        while os.path.exists(target):
            n += 1
            target = f"{base}.{n}"
        shutil.copy2(input_path, target)
        return target
    except OSError:
        return None


def _paragraph_text(para):
    return (para.text or "").strip()


def _detect_heading_level(para):
    style_name = ""
    try:
        style_name = para.style.name or ""
    except Exception:
        style_name = ""
    match = re.search(r"Heading\s+([1-4])", style_name, re.I)
    if match:
        return int(match.group(1))

    text = _paragraph_text(para)
    for level, pattern in HEADING_PATTERNS.items():
        if pattern.search(text):
            return level
    return 0


def _has_header_footer(doc):
    for section in doc.sections:
        if any((p.text or "").strip() for p in section.header.paragraphs):
            return True
        if section.footer.paragraphs:
            return True
    return False


def _has_complex_objects(doc):
    xml = doc._element.xml
    markers = ("<w:drawing", "<w:object", "<w:pict", "<w:sdt", "<w:ins", "<w:del")
    inline_count = xml.count("<wp:inline")
    drawing_count = xml.count("<w:drawing")
    if drawing_count > inline_count:
        return True
    return any(marker in xml for marker in markers if marker != "<w:drawing")


def _is_caption_para(text, prefixes):
    clean = (text or "").strip()
    for prefix in prefixes:
        if prefix and clean.startswith(prefix):
            return True
    return False


def _is_toc_style(para):
    """检测段落样式是否为目录条目样式（toc 1/2/3...）。"""
    try:
        name = para.style.name or ""
    except Exception:
        return False
    return name.lower().startswith("toc")


def _section_has_page_number(section):
    """检测 section 的页脚中是否已存在 PAGE 域。"""
    for instr in section.footer._element.iter(qn("w:instrText")):
        if instr.text and "PAGE" in instr.text:
            return True
    return False


def _build_body_section_maps(doc):
    """遍历 body 子元素，构建段落索引→section 索引、表格索引→section 索引映射。

    OOXML 中，分节符以段落 pPr 内的 sectPr 标记该段为当前节的最后一段；
    body 末尾的 sectPr 为最后一节的属性。
    """
    para_to_section = {}
    table_to_section = {}
    p_idx = t_idx = sec_idx = 0
    for child in doc.element.body.iterchildren():
        tag = child.tag.split("}", 1)[-1]
        if tag == "p":
            para_to_section[p_idx] = sec_idx
            ppr = child.find(qn("w:pPr"))
            if ppr is not None and ppr.find(qn("w:sectPr")) is not None:
                sec_idx += 1
            p_idx += 1
        elif tag == "tbl":
            table_to_section[t_idx] = sec_idx
            t_idx += 1
    return para_to_section, table_to_section


def _collect_cover_section_indices(doc):
    """返回封面/前置页 section 索引集合（需保留原版式）。

    判定：从首个 section 起连续无页码的 section 视为封面/前置页，直到遇到首个
    有页码的 section 为止；若全文均无页码，则不构成“封面/正文”区分，返回空集
    （按普通文档处理，正常规范化全部内容）。
    """
    has_pn = [_section_has_page_number(sec) for sec in doc.sections]
    if not any(has_pn):
        return set()
    cover = set()
    for i, has in enumerate(has_pn):
        if has:
            break
        cover.add(i)
    return cover


def _find_toc_title_index(doc):
    """找到首个 TOC 条目段落，向前回溯最近的非空非 TOC 段落索引（即“目录”标题段）。"""
    paragraphs = doc.paragraphs
    first_toc = None
    for i, p in enumerate(paragraphs):
        if _is_toc_style(p):
            first_toc = i
            break
    if first_toc is None:
        return None
    for j in range(first_toc - 1, -1, -1):
        p = paragraphs[j]
        if _is_toc_style(p):
            continue
        if _paragraph_text(p):
            return j
    return None


def _mark_toc_fields_dirty(doc):
    """将 TOC 域标记为 dirty，使 Word 打开时提示更新目录以同步页码与条目。"""
    begin_stack = []
    for el in doc.element.body.iter():
        tag = el.tag.split("}", 1)[-1]
        if tag == "fldChar":
            ftype = el.get(qn("w:fldCharType"))
            if ftype == "begin":
                begin_stack.append(el)
            elif ftype == "end" and begin_stack:
                begin_stack.pop()
        elif tag == "instrText":
            if el.text and "TOC" in el.text and begin_stack:
                begin_stack[-1].set(qn("w:dirty"), "true")


def _apply_run_defaults(para, zh_font, en_font, size_pt, bold=None, color=None):
    runs = para.runs
    if not runs and para.text:
        runs = [para.add_run()]
    for run in runs:
        set_run_font(run, zh_font, en_font, size_pt=size_pt, bold=bold, color=color)


def _apply_page_settings(doc, cfg):
    page = cfg.get("page", {})
    margins = page.get("margins", {})
    pn = cfg.get("pageNumber", {})
    only_existing = pn.get("onlyUpdateExisting", True)

    for sec in doc.sections:
        sec.page_width = Cm(21.0)
        sec.page_height = Cm(29.7)
        if "top" in margins:
            sec.top_margin = Cm(margins["top"])
        if "bottom" in margins:
            sec.bottom_margin = Cm(margins["bottom"])
        if "left" in margins:
            sec.left_margin = Cm(margins["left"])
        if "right" in margins:
            sec.right_margin = Cm(margins["right"])

        if not pn.get("enabled", True):
            continue
        # 封面/前置页：无既有页码则跳过，保留原始页脚（不注入页码）
        if only_existing and not _section_has_page_number(sec):
            continue

        footer = sec.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.clear()
        p.alignment = resolve_align(pn.get("align", "center"))
        set_para_spacing(p, line_spacing_mult=1.0, before_pt=0, after_pt=0)
        set_zero_indent(p)
        add_page_number_field(p, lambda r: set_run_font(
            r, pn.get("font", "宋体"), pn.get("asciiFont", "Times New Roman"),
            size_pt=resolve_size(pn.get("size", "小五"))))


def _apply_document_styles(doc, cfg):
    normal = doc.styles["Normal"]
    body = cfg.get("body", {})
    set_run_font_element(normal.element, body.get("font", "宋体"), body.get("asciiFont", "Times New Roman"),
                         resolve_size(body.get("size", "小四")), body.get("color", "000000"))
    normal.font.size = Pt(resolve_size(body.get("size", "小四")))
    npf = normal.paragraph_format
    npf.line_spacing = 1.0
    npf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    npf.first_line_indent = Pt(0)
    npf.space_before = Pt(0)
    npf.space_after = Pt(0)
    # 关键：清掉 Normal 样式 <w:ind> 里的字符单位属性（firstLineChars 等），
    # 否则单元格段落继承后会按字符数渲染首行缩进。
    clear_style_indent_chars(normal)

    for lv in range(1, 5):
        hc = cfg.get("headings", {}).get("h" + str(lv))
        if not hc:
            continue
        try:
            style = doc.styles["Heading " + str(lv)]
        except KeyError:
            continue
        set_run_font_element(style.element, hc.get("font", "黑体"), hc.get("asciiFont", "Times New Roman"),
                             resolve_size(hc.get("size", "小四")), hc.get("color", "000000"))
        style.font.size = Pt(resolve_size(hc.get("size", "小四")))
        style.font.bold = hc.get("bold", True)
        pf = style.paragraph_format
        pf.line_spacing = hc.get("lineSpacing", 1.5)
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.space_before = Pt(hc.get("spaceBeforePt", 0))
        pf.space_after = Pt(hc.get("spaceAfterPt", 0))


def _apply_heading_format(para, cfg, level):
    hc = cfg["headings"].get("h" + str(level), {})
    try:
        para.style = "Heading " + str(level)
    except Exception:
        pass
    para.alignment = resolve_align(hc.get("align", "left"))
    set_para_spacing(para, line_spacing_mult=hc.get("lineSpacing", 1.5),
                     before_pt=hc.get("spaceBeforePt", 0), after_pt=hc.get("spaceAfterPt", 0))
    set_zero_indent(para)
    _apply_run_defaults(
        para,
        hc.get("font", "黑体"),
        hc.get("asciiFont", cfg["body"].get("asciiFont", "Times New Roman")),
        resolve_size(hc.get("size", cfg["body"].get("size", "小四"))),
        bold=hc.get("bold", True),
        color=hc.get("color", "000000"),
    )


def _apply_body_format(para, cfg):
    body = cfg["body"]
    para.alignment = resolve_align(body.get("align", "justify"))
    set_para_spacing(para, line_spacing_mult=body.get("lineSpacing", 1.5),
                     before_pt=0, after_pt=body.get("spaceAfterPt", 0))
    indent_chars = body.get("firstLineIndentChars", 0)
    para.paragraph_format.first_line_indent = (
        Pt(resolve_size(body.get("size", "小四")) * indent_chars) if indent_chars > 0 else Pt(0)
    )
    para.paragraph_format.left_indent = Pt(0)
    para.paragraph_format.right_indent = Pt(0)
    _apply_run_defaults(
        para,
        body.get("font", "宋体"),
        body.get("asciiFont", "Times New Roman"),
        resolve_size(body.get("size", "小四")),
        color=body.get("color", "000000"),
    )


def _apply_caption_format(para, cfg, kind):
    c = cfg[kind]
    para.alignment = resolve_align(c.get("captionAlign", "center"))
    set_para_spacing(para, line_spacing_mult=c.get("captionLineSpacing", 1.0), before_pt=0, after_pt=2)
    set_zero_indent(para)
    _apply_run_defaults(
        para,
        c.get("captionText", cfg["body"].get("font", "宋体")),
        c.get("captionAsciiFont", cfg["body"].get("asciiFont", "Times New Roman")),
        resolve_size(c.get("captionSize", "五号")),
        bold=c.get("captionBold", True),
    )


def _apply_reference_format(para, cfg):
    r = cfg.get("references", {})
    para.alignment = resolve_align(r.get("align", "justify"))
    set_para_spacing(para, line_spacing_mult=r.get("lineSpacing", 1.0), before_pt=0, after_pt=3)
    hang_pt = r.get("hangingIndentPt", 24)
    para.paragraph_format.left_indent = Pt(hang_pt)
    para.paragraph_format.first_line_indent = Pt(-hang_pt)
    _apply_run_defaults(
        para,
        r.get("font", "宋体"),
        r.get("asciiFont", "Times New Roman"),
        resolve_size(r.get("size", "小五")),
    )


def _normalize_paragraphs(doc, cfg, para_to_section, cover_sections, on_warning=None):
    fig_prefix = cfg.get("figure", {}).get("prefix", "图")
    table_prefix = cfg.get("table", {}).get("prefix", "表")
    in_refs = False
    toc_title_index = _find_toc_title_index(doc)

    for idx, para in enumerate(doc.paragraphs):
        # 跳过目录条目段落，保留目录原格式（避免 TOC 文本被误判为标题）
        if _is_toc_style(para):
            continue
        # 跳过“目录”标题段，保留原格式
        if toc_title_index is not None and idx == toc_title_index:
            continue
        # 跳过封面/前置页 section 内的段落，保留封面版式
        sec_idx = para_to_section.get(idx)
        if sec_idx is not None and sec_idx in cover_sections:
            continue

        text = _paragraph_text(para)
        if not text:
            set_zero_indent(para)
            set_para_spacing(para, line_spacing_mult=1.0, before_pt=0, after_pt=0)
            continue

        if "参考文献" in text or "参考资料" in text:
            in_refs = True

        level = _detect_heading_level(para)
        if level:
            _apply_heading_format(para, cfg, level)
        elif _is_caption_para(text, (fig_prefix, "图", "Figure", "Fig")):
            _apply_caption_format(para, cfg, "figure")
        elif _is_caption_para(text, (table_prefix, "表", "Table")):
            _apply_caption_format(para, cfg, "table")
        elif in_refs and re.match(r"^\[\d+\]", text):
            _apply_reference_format(para, cfg)
        else:
            _apply_body_format(para, cfg)


def _normalize_tables(doc, cfg, table_to_section, cover_sections, on_warning=None):
    t = cfg.get("table", {})
    cell_font = t.get("cellFont", cfg["body"].get("font", "宋体"))
    cell_en = t.get("cellAsciiFont", cfg["body"].get("asciiFont", "Times New Roman"))
    cell_size = resolve_size(t.get("cellSize", cfg["body"].get("size", "小四")))
    cell_align = resolve_align(t.get("cellAlign", "center"))
    cell_valign = WD_ALIGN_VERTICAL.CENTER if t.get("cellVAlign", "center") == "center" else WD_ALIGN_VERTICAL.TOP

    for t_idx, tbl in enumerate(doc.tables):
        # 跳过封面/前置页 section 内的表格，保留封面表格原版式
        sec_idx = table_to_section.get(t_idx)
        if sec_idx is not None and sec_idx in cover_sections:
            continue
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for r_idx, row in enumerate(tbl.rows):
            for cell in row.cells:
                cell.vertical_alignment = cell_valign
                for para in cell.paragraphs:
                    para.alignment = cell_align
                    set_para_spacing(para, line_spacing_mult=t.get("cellLineSpacing", 1.0), before_pt=0, after_pt=0)
                    set_zero_indent(para)
                    _apply_run_defaults(
                        para,
                        cell_font,
                        cell_en,
                        cell_size,
                        bold=t.get("headerBold", True) if r_idx == 0 else False,
                    )


def _normalize_inline_images(doc, cfg, on_warning=None):
    figure = cfg.get("figure", {})
    max_width = figure.get("maxWidthCm")
    max_height = figure.get("maxHeightCm")
    if not max_width and not max_height:
        return

    max_w_emu = int(max_width * EMU_PER_CM) if max_width else None
    max_h_emu = int(max_height * EMU_PER_CM) if max_height else None
    for shape in getattr(doc, "inline_shapes", []):
        try:
            width = int(shape.width)
            height = int(shape.height)
        except Exception:
            continue
        if width <= 0 or height <= 0:
            continue
        scale = 1.0
        if max_w_emu and width > max_w_emu:
            scale = min(scale, max_w_emu / width)
        if max_h_emu and height > max_h_emu:
            scale = min(scale, max_h_emu / height)
        if scale < 1.0:
            shape.width = int(width * scale)
            shape.height = int(height * scale)
