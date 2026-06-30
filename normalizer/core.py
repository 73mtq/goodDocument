"""normalizer 主入口 + 路径校验 + 备份。"""
from __future__ import annotations

import os
import time
import zipfile
from typing import Optional

import lxml.etree

from docx import Document

from formatter import (
    resolve_size,
    resolve_align,
    set_run_font,
    set_para_spacing,
    set_zero_indent,
    clear_style_indent_chars,
    add_page_number_field,
    set_run_font_element,
)

from .errors import (
    NormalizeError,
    InputNotFoundError,
    InvalidFileTypeError,
    CorruptDocxError,
    OutputNotWritableError,
    SameInputOutputError,
    WarningCollector,
    NormalizeResult,
)
from .normalize import (
    HEADING_PATTERNS,
    _apply_run_defaults,
    _apply_page_settings,
    _apply_document_styles,
    _apply_heading_format,
    _apply_body_format,
    _apply_caption_format,
    _apply_reference_format,
    _normalize_paragraphs,
    _normalize_tables,
    _normalize_inline_images,
    _mark_toc_fields_dirty,
    _is_caption_para,
    _is_toc_style,
    _paragraph_text,
    _detect_heading_level,
    _find_toc_title_index,
    _build_body_section_maps,
    _collect_cover_section_indices,
    _section_has_page_number,
    _has_header_footer,
    _has_complex_objects,
    DocxInspection,
    inspect_docx,
)
from .paths import _same_path, _is_docx_path, _validate_paths, _backup_source


def normalize_docx(config, input_path, output_path, assets_dir=None, *,
                   backup=True, dry_run=False, return_result=False, on_warning=None):
    """规范化 docx。薄包装层。

    旧签名（不传新参数）行为完全兼容：返回 str(output_path)。
    新签名支持 backup / dry_run / return_result / on_warning。
    抛 NormalizeError 子类替代原本的 ValueError / FileNotFoundError / KeyError。
    """
    del assets_dir  # Kept for a stable public interface.
    start = time.time()
    input_path = os.fspath(input_path)
    output_path = os.fspath(output_path)

    # 1. 路径校验
    _validate_paths(input_path, output_path)

    # 2. 备份（dry_run 时跳过；备份失败归为 warning）
    backup_path: Optional[str] = None
    if backup and not dry_run:
        backup_path = _backup_source(input_path)
        if backup_path is None and on_warning:
            on_warning("备份失败（不影响规范化继续）", location="备份")

    # 3. 加载 + 规范化（try/except 把异常转成 NormalizeError）
    with WarningCollector() as wc:
        try:
            doc = Document(input_path)
        except (zipfile.BadZipFile, lxml.etree.XMLSyntaxError) as e:
            raise CorruptDocxError(
                f"docx 文件已损坏: {input_path}",
                cause=e, hint="用 Word 打开重新保存一次",
            ) from e
        except KeyError as e:
            raise CorruptDocxError(
                f"docx 关键样式缺失: {input_path}",
                cause=e, hint="用 Word 打开重新保存一次",
            ) from e
        except Exception as e:  # noqa: BLE001
            type_name = type(e).__name__
            if "Package" in type_name or "Zip" in type_name or "Found" in type_name:
                raise CorruptDocxError(
                    f"docx 文件已损坏: {input_path}",
                    cause=e, hint="用 Word 打开重新保存一次",
                ) from e
            raise

        try:
            para_to_section, table_to_section = _build_body_section_maps(doc)
            cover_sections = _collect_cover_section_indices(doc)

            _apply_page_settings(doc, config)
            _apply_document_styles(doc, config)
            _normalize_paragraphs(doc, config, para_to_section, cover_sections, on_warning=wc.warn)
            _normalize_tables(doc, config, table_to_section, cover_sections, on_warning=wc.warn)
            _normalize_inline_images(doc, config, on_warning=wc.warn)
            _mark_toc_fields_dirty(doc)
        except PermissionError as e:
            raise OutputNotWritableError(
                f"无法访问文件: {output_path}",
                cause=e, hint="检查目录权限和磁盘空间",
            ) from e

        # 4. 写文件（dry_run 时跳过）
        if not dry_run:
            try:
                out_dir = os.path.dirname(os.path.abspath(output_path))
                if out_dir:
                    os.makedirs(out_dir, exist_ok=True)
                doc.save(output_path)
            except PermissionError as e:
                raise OutputNotWritableError(
                    f"无法写入输出: {output_path}",
                    cause=e, hint="检查目录权限和磁盘空间",
                ) from e
            except OSError as e:
                raise OutputNotWritableError(
                    f"无法写入输出: {output_path}",
                    cause=e, hint="检查目录权限和磁盘空间",
                ) from e

    duration_ms = int((time.time() - start) * 1000)
    if return_result:
        return NormalizeResult(
            input_path=input_path,
            output_path=output_path if not dry_run else None,
            backup_path=backup_path,
            dry_run=dry_run,
            paragraphs_processed=len(doc.paragraphs),
            tables_processed=len(doc.tables),
            images_processed=len(getattr(doc, "inline_shapes", [])),
            changes=[],
            warnings=list(wc.warnings),
            errors=[],
            duration_ms=duration_ms,
        )
    return output_path
